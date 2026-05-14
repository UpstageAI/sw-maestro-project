"""Build eval fixtures deterministically.

Generates:
  eval/forms/{bk21,nrf_undergrad,conference_support}.hwpx
  eval/form_labels.json — ground-truth items per form
  eval/materials/{학부생,석사1,석사2}/cv.txt          — anonymised CV text
  eval/materials/{학부생,석사1,석사2}/plan.docx        — research plan with tables
  eval/materials/{학부생,석사1,석사2}/report.pdf       — research notes with tables

Regenerate: `uv run python eval/build_eval.py`
"""

from __future__ import annotations

import io
import json
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path

ROOT = Path(__file__).parent
FORMS_DIR = ROOT / "forms"
MATERIALS_DIR = ROOT / "materials"
LABELS_PATH = ROOT / "form_labels.json"

MIMETYPE = b"application/hwp+zip"

CONTAINER_XML = b"""\
<?xml version="1.0" encoding="UTF-8"?>
<container xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="Contents/content.hpf" media-type="application/hwp+zip"/>
  </rootfiles>
</container>
"""

CONTENT_HPF = b"""\
<?xml version="1.0" encoding="UTF-8"?>
<hpf:HWPFDocumentInformation
    xmlns:hpf="http://www.hancom.co.kr/hwpml/2011/hp">
  <hpf:manifest>
    <hpf:item id="Contents/section0.xml" mediaType="application/xml"/>
  </hpf:manifest>
</hpf:HWPFDocumentInformation>
"""


@dataclass
class FormSpec:
    file_id: str  # filename stem
    title: str  # human-readable
    paragraphs: list[tuple[str, bool]]  # (label, is_pii) — body paragraphs
    table_headers: list[str]  # 1 row of headers + 1 empty row
    pii_labels: frozenset[str]  # exact labels that are PII (parser strips prefixes via [:40])


def _section_xml(spec: FormSpec) -> bytes:
    para_blocks = "\n".join(
        f"  <hp:p><hp:run><hp:t>{label}</hp:t></hp:run></hp:p>" for label, _ in spec.paragraphs
    )
    header_cells = "\n      ".join(
        f"<hp:tc><hp:subList><hp:p><hp:run><hp:t>{h}</hp:t></hp:run></hp:p></hp:subList></hp:tc>"
        for h in spec.table_headers
    )
    empty_cells = "\n      ".join(
        "<hp:tc><hp:subList><hp:p><hp:run><hp:t></hp:t></hp:run></hp:p></hp:subList></hp:tc>"
        for _ in spec.table_headers
    )
    return f"""\
<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section"
        xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
{para_blocks}
  <hp:tbl>
    <hp:tr>
      {header_cells}
    </hp:tr>
    <hp:tr>
      {empty_cells}
    </hp:tr>
  </hp:tbl>
</hs:sec>
""".encode()


def _write_hwpx(spec: FormSpec) -> Path:
    out = FORMS_DIR / f"{spec.file_id}.hwpx"
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
        info = zipfile.ZipInfo("mimetype")
        info.compress_type = zipfile.ZIP_STORED
        z.writestr(info, MIMETYPE)
        z.writestr("META-INF/container.xml", CONTAINER_XML)
        z.writestr("Contents/content.hpf", CONTENT_HPF)
        z.writestr("Contents/section0.xml", _section_xml(spec))
    out.write_bytes(buf.getvalue())
    return out


def _labels_from_parser(spec: FormSpec, hwpx_bytes: bytes) -> dict:
    """Derive labels by parsing the built .hwpx — guarantees label set matches parser."""
    from backend.app.hwpx.parser import parse_hwpx

    doc = parse_hwpx(hwpx_bytes)
    items = []
    for it in doc.items:
        items.append(
            {
                "item_id": it.item_id,
                "label": it.label,
                "section": it.section,
                "expected_chars": 200,
                "is_pii": it.label in spec.pii_labels,
            }
        )
    return {
        "file_id": spec.file_id,
        "title": spec.title,
        "items": items,
        "table_headers": spec.table_headers,
    }


# ---------------------------------------------------------------------------
# Form specs
# ---------------------------------------------------------------------------

FORMS: list[FormSpec] = [
    FormSpec(
        file_id="bk21",
        title="BK21 사업 신청서",
        paragraphs=[
            ("신청자 성명", True),
            ("소속 대학교", False),
            ("연구의 필요성", False),
            ("선행 연구 분석", False),
            ("연구 방법", False),
            ("예상 결과 및 활용 방안", False),
        ],
        table_headers=["분기", "추진 내용", "비고"],
        pii_labels=frozenset({"신청자 성명"}),
    ),
    FormSpec(
        file_id="nrf_undergrad",
        title="한국연구재단 학부생연구지원 신청서",
        paragraphs=[
            ("학번", True),
            ("이메일", True),
            ("연락처", True),
            ("연구 주제", False),
            ("연구 배경", False),
            ("연구 목표", False),
            ("연구 일정", False),
        ],
        table_headers=["항목", "금액", "사용 사유"],
        pii_labels=frozenset({"학번", "이메일", "연락처"}),
    ),
    FormSpec(
        file_id="conference_support",
        title="학회 발표 지원 신청서",
        paragraphs=[
            ("성명", True),
            ("주민등록번호", True),
            ("계좌번호", True),
            ("발표 논문 제목", False),
            ("학회명", False),
            ("발표 일자", False),
            ("발표 초록", False),
        ],
        table_headers=["공동저자 성명", "소속", "역할"],
        pii_labels=frozenset({"성명", "주민등록번호", "계좌번호", "공동저자 성명"}),
    ),
]


# ---------------------------------------------------------------------------
# Mock materials (no real PII) — CV (txt) + plan (docx) + report (pdf)
# ---------------------------------------------------------------------------


_FONT_CANDIDATES = [
    Path("C:/Windows/Fonts/malgun.ttf"),
    Path("C:/Windows/Fonts/NanumGothic.ttf"),
    Path("/usr/share/fonts/truetype/nanum/NanumGothic.ttf"),
    Path("/System/Library/Fonts/AppleGothic.ttf"),
]


def _korean_font() -> Path:
    for p in _FONT_CANDIDATES:
        if p.exists():
            return p
    raise FileNotFoundError(
        "Korean TTF not found; install one of: " + ", ".join(str(p) for p in _FONT_CANDIDATES)
    )


@dataclass
class PersonaMaterials:
    persona: str
    cv_text: str
    plan_title: str
    plan_summary: str
    plan_budget_rows: list[tuple[str, str, str]]  # (항목, 금액(만원), 사용 사유)
    plan_schedule_rows: list[tuple[str, str, str]]  # (분기, 마일스톤, 산출물)
    report_title: str
    report_intro: str
    report_results_rows: list[tuple[str, str, str, str]]  # (모델, 데이터, 지표, 점수)
    report_comparison_rows: list[tuple[str, str, str]]  # (조건, 정확도, 비고)


PERSONAS: list[PersonaMaterials] = [
    PersonaMaterials(
        persona="학부생",
        cv_text=(
            "[익명 이력서 — 신원 식별 정보 제거됨]\n"
            "전공: 컴퓨터공학과 (학부 4학년)\n"
            "관심 분야: 한국어 자연어처리, 학술 글쓰기 지원 LLM, 도메인 특화 미세조정\n\n"
            "## 주요 프로젝트\n"
            "- 한국어 문법 교정 모듈 (캡스톤 디자인, 2025)\n"
            "  · 8,000개 문장 데이터셋, BERT 기반 분류기 + 규칙 후처리\n"
            "  · 평가 정확도 87.2%, 교내 캡스톤 우수상 수상\n"
            "- BERT 기반 한국어 감성 분석 (개인 프로젝트, 2024)\n"
            "  · NSMC + 직접 수집 도메인 데이터로 fine-tune\n"
            "- ATIS 한국어 의도 분류 (2024 가을 학기 수업)\n\n"
            "## 프로젝트별 정확도 (테스트셋, %)\n"
            "한국어 문법 교정     ████████████████████  87.2\n"
            "감성 분석 (NSMC)     ███████████████░░░░░  72.1\n"
            "ATIS 의도 분류       ██████████████░░░░░░  68.5\n\n"
            "## 기술 스택\n"
            "- Python (PyTorch / HuggingFace), 한국어 토크나이저 (Mecab-ko, KoNLPy)\n"
            "- 데이터 라벨링 도구 자체 구현 경험 (Streamlit 기반)\n"
            "- Docker / GitHub Actions 기본 사용 경험\n\n"
            "## 발표 / 수상\n"
            "- 교내 캡스톤 디자인 우수상 (2025)\n"
            "- 학부생 학술 콜로키움 포스터 발표 (2025)\n"
        ),
        plan_title="한국어 학술 글쓰기 지원을 위한 LLM 기반 문장 다듬기 도구 개발",
        plan_summary=(
            "비전공 학부생이 작성한 한국어 학술 텍스트는 어휘·문법 오류는 적어도 "
            "학술 문어체로서의 일관성과 논리 흐름이 부족한 경우가 많다. 본 연구는 "
            "Upstage Solar 모델을 기반으로 도메인 특화 프롬프트와 소규모 RLHF 데이터를 "
            "활용하여, 학부생 보고서 수준의 한국어 텍스트를 학술 문어체로 다듬는 도구를 "
            "개발한다. 1년차에는 8,000개 문장 데이터셋 구축과 baseline 모델 평가를, "
            "2년차에는 RLHF fine-tuning과 사용자 평가를 진행한다."
        ),
        plan_budget_rows=[
            ("인건비 (학부 연구원 1인)", "600", "월 50만원 × 12개월"),
            ("데이터 라벨링", "300", "전문 첨삭 평가자 200문장 × 6회"),
            ("API / GPU 사용료", "240", "Solar API 호출 + RTX 4090 임대"),
            ("학회 참가비", "120", "정보과학회 동계 워크숍 1회"),
            ("도서 / 출판", "60", "참고 문헌 / 결과 정리 인쇄"),
            ("총계", "1,320", "1년차 예산"),
        ],
        plan_schedule_rows=[
            ("1분기", "데이터 수집·정제", "8,000문장 정제 데이터셋 v1"),
            ("2분기", "Baseline 평가", "BLEU·BERTScore 기준 표"),
            ("3분기", "RLHF 데이터 수집", "선호도 라벨 1,200쌍"),
            ("4분기", "Fine-tune·사용자 평가", "사용자 만족도 설문 결과"),
        ],
        report_title="문장 다듬기 도구 v1 — 베이스라인 평가 결과",
        report_intro=(
            "본 보고서는 연구계획서의 1분기·2분기 산출물에 해당하며, 8,000문장 "
            "정제 데이터셋(원문/모범문 쌍)에 대한 baseline 모델 평가 결과를 정리한 것이다. "
            "평가는 BLEU-4, BERTScore-F1, 그리고 학술 문어체 적합도 사람 평가(5점 척도) "
            "세 지표로 수행하였다."
        ),
        report_results_rows=[
            ("모델", "데이터셋", "지표", "점수"),
            ("Solar-pro (zero-shot)", "in-house v1", "BLEU-4", "22.3"),
            ("Solar-pro (zero-shot)", "in-house v1", "BERTScore-F1", "0.78"),
            ("Solar-pro (few-shot 8)", "in-house v1", "BLEU-4", "26.1"),
            ("Solar-pro (few-shot 8)", "in-house v1", "BERTScore-F1", "0.82"),
            ("Solar-pro + 규칙 후처리", "in-house v1", "BLEU-4", "28.7"),
            ("Solar-pro + 규칙 후처리", "in-house v1", "사람 평가", "4.1 / 5.0"),
        ],
        report_comparison_rows=[
            ("Zero-shot vs Few-shot", "+3.8 BLEU", "few-shot 예시 8개로 의미 향상"),
            ("Few-shot vs 후처리 추가", "+2.6 BLEU", "조사·어미 정규화 규칙 효과"),
            ("Baseline vs 최종", "+6.4 BLEU", "전체 파이프라인 기여"),
        ],
    ),
    PersonaMaterials(
        persona="석사1",
        cv_text=(
            "[익명 이력서 — 신원 식별 정보 제거됨]\n"
            "전공: 데이터사이언스 (석사 1학기)\n"
            "관심 분야: 헬스케어 NLP, 한국어 임상 노트 요약, ClinicalBERT-Ko\n\n"
            "## 학술 활동\n"
            "- 한국정보과학회 2025 동계 워크숍 포스터 발표\n"
            '  · 제목: "한국어 임상 노트의 추출형 요약을 위한 도메인 적응 BERT 분석"\n'
            '- 학부 졸업논문(2024): "의료 도메인 토큰화 전략 비교 연구"\n\n'
            "## 보유 경험\n"
            "- 임상 텍스트 사전처리 파이프라인(자체 구현, 2024) — 의료 약어 정규화 모듈 포함\n"
            "- ClinicalBERT-Ko 사전훈련 모델 fine-tune 경험 (3개 task)\n"
            "- 의료 IRB 협력 절차 문서 작성 경험 (학부 인턴 기간 보조)\n\n"
            "## 사전 실험 결과 — ROUGE-L 분포 (참고용)\n"
            "BERT-base-multilingual    ███████████░░░░  0.34\n"
            "KoBERT                    █████████████░░  0.39\n"
            "ClinicalBERT-Ko (proposed) ██████████████░  0.41\n\n"
            "## 기술 스택\n"
            "- Python / PyTorch, HuggingFace Transformers, NLTK Korean fork\n"
            "- 의료 데이터 비식별화 도구 운용 경험 (Presidio Korean 커스텀 인식기)\n"
        ),
        plan_title="한국어 임상 노트 자동 요약을 통한 의료진 업무 부담 완화",
        plan_summary=(
            "한국어 임상 노트는 영문 대비 자연어처리 자원과 사전훈련 모델이 부족하여, "
            "의료진이 환자 인계·교대 시 매번 노트 전체를 통독해야 하는 비효율이 있다. "
            "본 연구는 ClinicalBERT-Ko 기반 추출형·생성형 하이브리드 요약 모델을 개발하여 "
            "노트 검토 시간을 30% 단축하는 것을 목표로 한다. 협력 병원 IRB 승인 후 "
            "비식별화된 노트 50,000건을 학습·평가 데이터로 활용한다."
        ),
        plan_budget_rows=[
            ("인건비 (석사 연구원)", "1,800", "월 150만원 × 12개월"),
            ("협력 병원 IRB 행정비", "200", "IRB 신청·심의 비용"),
            ("데이터 비식별화 도구", "300", "Presidio + 자체 NER 라이선스"),
            ("GPU / 클라우드", "600", "A100 인스턴스 800시간"),
            ("의료진 평가 사례비", "400", "현직 의료진 10인 × 5건 × 8만원"),
            ("학회 참가 / 출장", "200", "MICCAI / 정보과학회"),
            ("총계", "3,500", "1년차 예산"),
        ],
        plan_schedule_rows=[
            ("1분기", "IRB 승인·데이터 비식별화 파이프라인", "비식별화 50,000건"),
            ("2분기", "ClinicalBERT-Ko fine-tune", "추출형 요약 baseline"),
            ("3분기", "생성형 요약 결합 / RLHF", "하이브리드 모델 v1"),
            ("4분기", "임상 사용자 평가", "노트 검토 시간 단축 보고"),
        ],
        report_title="ClinicalBERT-Ko 베이스라인 평가 — 1차 결과",
        report_intro=(
            "본 보고서는 연구계획서의 2분기 산출물로, ClinicalBERT-Ko 기반 추출형 요약 "
            "모델의 사전 실험 결과를 정리한다. 비식별화된 임상 노트 8,200건(연구계획상 "
            "최종 50,000건의 일부)에 대해 ROUGE-L과 의료진 정성 평가 두 축으로 측정하였다."
        ),
        report_results_rows=[
            ("모델", "데이터셋", "지표", "점수"),
            ("BERT-multilingual", "임상노트 v0.2", "ROUGE-L", "0.34"),
            ("KoBERT", "임상노트 v0.2", "ROUGE-L", "0.39"),
            ("ClinicalBERT-Ko", "임상노트 v0.2", "ROUGE-L", "0.41"),
            ("ClinicalBERT-Ko + LLM 후처리", "임상노트 v0.2", "ROUGE-L", "0.46"),
            ("ClinicalBERT-Ko + LLM 후처리", "임상노트 v0.2", "의료진 만족도", "4.0 / 5.0"),
        ],
        report_comparison_rows=[
            ("multilingual vs 도메인 특화", "+0.07 ROUGE-L", "도메인 사전훈련 효과 확인"),
            ("추출형 vs 하이브리드", "+0.05 ROUGE-L", "LLM 후처리로 자연성 향상"),
            ("자동 vs 의료진 평가", "—", "ROUGE-L과 만족도 상관 0.62"),
        ],
    ),
    PersonaMaterials(
        persona="석사2",
        cv_text=(
            "[익명 이력서 — 신원 식별 정보 제거됨]\n"
            "전공: 인공지능 (석사 3학기)\n"
            "관심 분야: 멀티모달 학습, 영상-텍스트 결합 추론, 한국어 OOD 견고성\n\n"
            "## 학술 활동\n"
            "- 국내 학술지 게재 1편 (한국정보과학회논문지, 2025): "
            '"한국어 멀티모달 모델의 분포 외 입력에 대한 견고성 분석"\n'
            "- 국제 컨퍼런스 워크숍 발표 1편 (NeurIPS Multimodal Workshop, 2024)\n\n"
            "## 보유 경험\n"
            "- CLIP·BLIP·OFA 등 4개 멀티모달 모델 패밀리 평가 파이프라인 구축\n"
            "- 한국어 OOD 벤치마크 자체 구성 (1차 5,000샘플)\n"
            "- 도메인 적응(domain adaptation) 기법 비교 분석\n\n"
            "## 4종 모델 OOD 정확도 (1차 5,000샘플 기준)\n"
            "CLIP-large                ██████████████░░░░  68.2\n"
            "BLIP-2                    ████████████████░░  72.5\n"
            "OFA-base                  ███████████░░░░░░░  62.1\n"
            "Ko-CLIP (proposed)        ██████████████████  78.4\n\n"
            "## 기술 스택\n"
            "- PyTorch / Transformers, OpenCLIP, MMEngine\n"
            "- 분산 학습(8 GPU), 대규모 데이터 파이프라인 구축\n"
        ),
        plan_title="영상-텍스트 결합 모델의 한국어 OOD 견고성 분석 및 도메인 적응",
        plan_summary=(
            "다국어 멀티모달 모델은 영문 벤치마크에서는 우수하나, 한국어 입력의 분포 외 "
            "(out-of-distribution; OOD) 케이스에서는 정확도 저하가 두드러진다. 본 연구는 "
            "한국어 OOD 벤치마크를 5,000→20,000샘플로 확장하고, 4종 멀티모달 모델 패밀리에 "
            "대해 도메인 적응 기법(가중치 fine-tune·prompt tuning·adapter)을 비교 분석한다. "
            "최종 산출물로 한국어 멀티모달 평가 표준을 제안한다."
        ),
        plan_budget_rows=[
            ("인건비 (석사 연구원)", "1,800", "월 150만원 × 12개월"),
            ("데이터 라벨링 / 검수", "500", "한국어 OOD 15,000샘플 추가"),
            ("GPU / 클라우드", "1,500", "A100 8-way × 2,000시간"),
            ("국제 학회 출장", "400", "NeurIPS / CVPR 발표 비용"),
            ("도메인 적응 라이선스", "100", "어댑터 라이브러리 / 평가 도구"),
            ("총계", "4,300", "1년차 예산"),
        ],
        plan_schedule_rows=[
            ("1분기", "OOD 벤치마크 확장", "한국어 20,000샘플 v2"),
            ("2분기", "4종 모델 baseline 평가", "모델별 정확도 표"),
            ("3분기", "도메인 적응 비교", "기법별 정확도 향상량"),
            ("4분기", "표준 제안 / 논문 작성", "국제 학회 투고 1편"),
        ],
        report_title="한국어 OOD 1차 평가 — 4종 멀티모달 모델 비교",
        report_intro=(
            "본 보고서는 연구계획서의 1·2분기 진행 결과로, 한국어 OOD 벤치마크 1차 "
            "5,000샘플에 대해 4종 멀티모달 모델 패밀리(CLIP-large, BLIP-2, OFA-base, "
            "Ko-CLIP-proposed)의 정확도를 측정한 결과를 정리한다. 평가 지표는 Top-1 분류 "
            "정확도와 영문 분포 대비 정확도 차이(domain gap) 두 가지이다."
        ),
        report_results_rows=[
            ("모델", "데이터셋", "지표", "점수"),
            ("CLIP-large", "한국어 OOD v1 (5K)", "Top-1 정확도", "68.2"),
            ("BLIP-2", "한국어 OOD v1 (5K)", "Top-1 정확도", "72.5"),
            ("OFA-base", "한국어 OOD v1 (5K)", "Top-1 정확도", "62.1"),
            ("Ko-CLIP (proposed)", "한국어 OOD v1 (5K)", "Top-1 정확도", "78.4"),
            ("Ko-CLIP (proposed)", "한국어 OOD v1 (5K)", "영문 대비 gap", "-7.6%p"),
        ],
        report_comparison_rows=[
            ("CLIP vs Ko-CLIP", "+10.2%p", "한국어 사전훈련 효과 검증"),
            ("BLIP-2 vs Ko-CLIP", "+5.9%p", "영상-텍스트 결합 구조 영향"),
            ("OFA vs Ko-CLIP", "+16.3%p", "OFA는 한국어 OOD에 가장 취약"),
        ],
    ),
]


def _write_cv_txt(out_dir: Path, p: PersonaMaterials) -> None:
    (out_dir / "cv.txt").write_text(p.cv_text, encoding="utf-8")


def _write_plan_docx(out_dir: Path, p: PersonaMaterials) -> None:
    import docx

    doc = docx.Document()
    doc.add_heading(p.plan_title, level=1)

    doc.add_heading("1. 연구 개요", level=2)
    doc.add_paragraph(p.plan_summary)

    doc.add_heading("2. 예산 계획 (단위: 만원)", level=2)
    bt = doc.add_table(rows=1 + len(p.plan_budget_rows), cols=3)
    bt.style = "Light Grid"
    for i, hdr in enumerate(("항목", "금액", "사용 사유")):
        bt.cell(0, i).text = hdr
    for r, row in enumerate(p.plan_budget_rows, start=1):
        for c, val in enumerate(row):
            bt.cell(r, c).text = val

    doc.add_heading("3. 추진 일정", level=2)
    st = doc.add_table(rows=1 + len(p.plan_schedule_rows), cols=3)
    st.style = "Light Grid"
    for i, hdr in enumerate(("분기", "마일스톤", "산출물")):
        st.cell(0, i).text = hdr
    for r, row in enumerate(p.plan_schedule_rows, start=1):
        for c, val in enumerate(row):
            st.cell(r, c).text = val

    doc.save(out_dir / "plan.docx")


def _write_report_pdf(out_dir: Path, p: PersonaMaterials) -> None:
    from fpdf import FPDF

    pdf = FPDF()
    font = _korean_font()
    pdf.add_font("Korean", "", str(font))
    pdf.set_font("Korean", size=11)

    pdf.add_page()
    pdf.set_font("Korean", size=14)
    pdf.cell(0, 10, p.report_title)
    pdf.ln(12)

    pdf.set_font("Korean", size=11)
    pdf.multi_cell(0, 6, p.report_intro)
    pdf.ln(4)

    pdf.set_font("Korean", size=12)
    pdf.cell(0, 8, "■ 표 1. 모델별 평가 결과")
    pdf.ln(8)
    pdf.set_font("Korean", size=10)
    col_w = [50, 50, 40, 40]
    for row in p.report_results_rows:
        for i, val in enumerate(row):
            pdf.cell(col_w[i], 7, val[: int(col_w[i] / 3)], border=1)
        pdf.ln(7)
    pdf.ln(4)

    pdf.set_font("Korean", size=12)
    pdf.cell(0, 8, "■ 표 2. 조건별 정확도 차이")
    pdf.ln(8)
    pdf.set_font("Korean", size=10)
    col_w2 = [60, 40, 80]
    pdf.cell(col_w2[0], 7, "조건", border=1)
    pdf.cell(col_w2[1], 7, "지표 차이", border=1)
    pdf.cell(col_w2[2], 7, "비고", border=1)
    pdf.ln(7)
    for row in p.report_comparison_rows:
        for i, val in enumerate(row):
            pdf.cell(col_w2[i], 7, val[: int(col_w2[i] / 3)], border=1)
        pdf.ln(7)

    pdf.output(str(out_dir / "report.pdf"))


def _write_persona_materials(p: PersonaMaterials) -> None:
    out_dir = MATERIALS_DIR / p.persona
    out_dir.mkdir(parents=True, exist_ok=True)
    # Wipe legacy plan.txt / report.txt so the directory matches the manifest above.
    for legacy in ("plan.txt", "report.txt"):
        legacy_path = out_dir / legacy
        if legacy_path.exists():
            legacy_path.unlink()
    _write_cv_txt(out_dir, p)
    _write_plan_docx(out_dir, p)
    _write_report_pdf(out_dir, p)
    print(f"wrote materials for {p.persona}")


def main() -> None:
    FORMS_DIR.mkdir(parents=True, exist_ok=True)
    MATERIALS_DIR.mkdir(parents=True, exist_ok=True)

    labels_out: list[dict] = []
    for spec in FORMS:
        path = _write_hwpx(spec)
        print(f"wrote {path} ({path.stat().st_size} bytes)")
        labels_out.append(_labels_from_parser(spec, path.read_bytes()))

    LABELS_PATH.write_text(json.dumps(labels_out, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"wrote {LABELS_PATH}")

    try:
        for p in PERSONAS:
            _write_persona_materials(p)
    except FileNotFoundError as exc:
        print(f"materials skipped — {exc}", file=sys.stderr)
        return


if __name__ == "__main__":
    main()
