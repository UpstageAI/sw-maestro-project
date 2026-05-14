"""
Build material test fixtures idempotently.

PDF strategy: fpdf2 (dev dependency) with Malgun Gothic (Windows) or NanumGothic
as a Korean-capable TrueType font. fpdf2 embeds the font subset so pypdf and
pdfplumber can extract Unicode text via the ToUnicode CMap.

Run:  py backend/tests/fixtures/materials/_build_materials.py
"""

from pathlib import Path

import docx
from fpdf import FPDF

OUT = Path(__file__).parent

# ---------------------------------------------------------------------------
# Korean font discovery
# ---------------------------------------------------------------------------

_FONT_CANDIDATES = [
    Path("C:/Windows/Fonts/malgun.ttf"),  # Windows: Malgun Gothic
    Path("C:/Windows/Fonts/NanumGothic.ttf"),  # Windows: Nanum Gothic
    Path("/usr/share/fonts/truetype/nanum/NanumGothic.ttf"),  # Ubuntu: Nanum
    Path("/System/Library/Fonts/AppleGothic.ttf"),  # macOS
]


def _find_korean_font() -> Path:
    for p in _FONT_CANDIDATES:
        if p.exists():
            return p
    raise FileNotFoundError(
        "No Korean TTF font found. Install one of: " + ", ".join(str(p) for p in _FONT_CANDIDATES)
    )


# ---------------------------------------------------------------------------
# notes.txt
# ---------------------------------------------------------------------------


def build_txt() -> None:
    path = OUT / "notes.txt"
    path.write_text("연구 노트\n- 가설: A → B\n- 실험: ...\n", encoding="utf-8")
    print(f"Written {path} ({path.stat().st_size} bytes)")


# ---------------------------------------------------------------------------
# plan.docx
# ---------------------------------------------------------------------------


def build_docx() -> None:
    path = OUT / "plan.docx"
    doc = docx.Document()
    doc.add_paragraph("연구 계획")
    doc.add_paragraph("1년차에는 데이터셋 구축에 집중한다.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "단계"
    table.cell(0, 1).text = "목표"
    table.cell(1, 0).text = "1단계"
    table.cell(1, 1).text = "데이터 수집"
    doc.save(path)
    print(f"Written {path} ({path.stat().st_size} bytes)")


# ---------------------------------------------------------------------------
# cv.pdf
# ---------------------------------------------------------------------------


def build_pdf() -> None:
    font_path = _find_korean_font()

    pdf = FPDF()
    pdf.add_font("Korean", "", str(font_path))
    pdf.set_font("Korean", size=12)

    pdf.add_page()
    pdf.cell(0, 10, "이름: 김연구")
    pdf.ln()
    pdf.cell(0, 10, "학력: 서울대학교 전기공학과")

    pdf.add_page()
    pdf.cell(0, 10, "연구분야: 머신러닝")

    path = OUT / "cv.pdf"
    pdf.output(str(path))
    print(f"Written {path} ({path.stat().st_size} bytes)")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    build_txt()
    build_docx()
    build_pdf()
    print("All fixtures ready.")


if __name__ == "__main__":
    main()
