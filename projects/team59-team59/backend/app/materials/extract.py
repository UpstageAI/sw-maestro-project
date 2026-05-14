"""Material text extraction: PDF, docx, txt, hwpx."""

from pathlib import Path

import docx
import pdfplumber
import pypdf

from backend.app.hwpx.parser import parse_hwpx


def extract_text(path: Path | str) -> str:
    """Extract plain text from a material file (PDF, docx, txt, or hwpx)."""
    path = Path(path)
    ext = path.suffix.lstrip(".").lower()
    if ext == "pdf":
        return _extract_pdf(path)
    if ext == "docx":
        return _extract_docx(path)
    if ext == "txt":
        return _extract_txt(path)
    if ext == "hwpx":
        return _extract_hwpx(path)
    raise ValueError(f"Unsupported material format: .{ext}")


def _extract_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="cp949")


def _extract_docx(path: Path) -> str:
    doc = docx.Document(path)
    paragraphs = "\n".join(p.text for p in doc.paragraphs)
    table_blocks = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append("\t".join(cell.text for cell in row.cells))
        table_blocks.append("\n".join(rows))
    parts = [paragraphs] + table_blocks
    return "\n\n".join(p for p in parts if p)


def _extract_hwpx(path: Path) -> str:
    try:
        doc = parse_hwpx(path.read_bytes())
    except ValueError as e:
        raise ValueError(f"Failed to parse HWPX: {path.name}") from e
    lines = [item.label for item in doc.items]
    for table in doc.tables:
        if table.headers:
            lines.append("\t".join(table.headers))
    return "\n".join(lines)


def _extract_pdf(path: Path) -> str:
    try:
        with pypdf.PdfReader(path) as reader:
            if reader.is_encrypted:
                raise ValueError(f"PDF is encrypted: {path.name}")
            pages_text = [page.extract_text() or "" for page in reader.pages]
            text = "\n\n".join(pages_text)
    except (pypdf.errors.PdfStreamError, pypdf.errors.PdfReadError) as e:
        raise ValueError(f"Corrupt PDF: {path.name}") from e

    if text.strip():
        return text

    with pdfplumber.open(path) as pdf:
        page_parts = []
        for page in pdf.pages:
            parts = []
            page_text = page.extract_text() or ""
            if page_text:
                parts.append(page_text)
            for table in page.extract_tables() or []:
                rows = ["\t".join(cell or "" for cell in row) for row in table if row]
                if rows:
                    parts.append("[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")
            page_parts.append("\n".join(parts))
        return "\n\n".join(page_parts)
